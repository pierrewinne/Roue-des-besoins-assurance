#!/usr/bin/env python3
"""Generate Word report from the collegial synthesis on contract adequacy feature."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0, 13, 110)  # Baloise primary-700

NAVY = RGBColor(0, 13, 110)
DARK = RGBColor(30, 30, 30)
GREY = RGBColor(100, 100, 100)
GREEN = RGBColor(22, 135, 65)
AMBER = RGBColor(201, 118, 18)
RED = RGBColor(217, 48, 76)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def bullet(doc, text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.8)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(text)
        run2.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SYNTHESE COLLEGIALE')
run.font.size = Pt(28)
run.font.color.rgb = NAVY
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Fonctionnalite d'adequation contrats")
run.font.size = Pt(18)
run.font.color.rgb = NAVY

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Decision GO / NO-GO')
run.font.size = Pt(14)
run.font.color.rgb = GREY

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Roue des Besoins Assurance — Baloise Luxembourg')
run.font.size = Pt(11)
run.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Analyse collegiale de 14 agents specialises')
run.font.size = Pt(11)
run.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Mars 2026')
run.font.size = Pt(11)
run.font.color.rgb = GREY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════
doc.add_heading('Table des matieres', level=1)
toc_items = [
    '1. Verdict global',
    '2. Synthese par equipe',
    '   2.1 Team Produit & Strategie',
    '   2.2 Team Gouvernance & Conformite',
    '   2.3 Team Technique Assurance',
    '   2.4 Team IT & Securite',
    '   2.5 Team Ops',
    '   2.6 Team Creative',
    '3. Convergences majeures (14/14)',
    '4. Divergences et arbitrages',
    '5. Benchmark Extraction Automatique',
    '   5.1 Solutions evaluees (8)',
    '   5.2 Scoring consolide',
    '   5.3 Divergences inter-agents (7)',
    '   5.4 Recommandation PM : scenario retenu',
    '   5.5 Architecture cible',
    '   5.6 Roadmap extraction',
    '6. Pre-requis avant developpement (Gates)',
    '7. Estimation de charge',
    '8. Risques residuels',
    '9. Recommandation finale consolidee',
    'Annexe A : Standards du marche luxembourgeois',
    'Annexe B : Fichiers d\'implementation',
    'Annexe C : Protocole QA benchmark (223 cas de test)',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. VERDICT GLOBAL
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. Verdict global', level=1)

p = doc.add_paragraph()
run = p.add_run('VERDICT : GO CONDITIONNEL (unanime, 14/14 agents)')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = GREEN

doc.add_paragraph(
    "Aucun agent n'a emis d'avis NO-GO. Tous convergent vers un GO sous conditions, "
    "avec des pre-requis variant selon le domaine. Les 3 conditions bloquantes transverses sont :"
)

add_table(doc,
    ['#', 'Condition suspensive', 'Porte par'],
    [
        ['C1', "Qualification IDD tranchee (information vs conseil) avant toute ligne de code", 'Compliance + Legal'],
        ['C2', "DPIA realisee et validee (donnees sensibles, art. 9 RGPD pour sante)", 'Compliance + Security'],
        ['C3', "Phase 1 = upload automatise avec extraction intelligente (pas de saisie manuelle)", 'IT Architect + Product Manager'],
    ],
    col_widths=[1, 10, 4]
)

p = doc.add_paragraph()
run = p.add_run("Note strategique : ")
run.bold = True
p.add_run(
    "L'approche retenue est l'upload direct de documents (PDF, photo) avec extraction automatique. "
    "L'utilisateur ne fait qu'uploader ses contrats — le systeme se charge de l'analyse. "
    "Simplicite maximale cote utilisateur, complexite geree cote serveur."
)

# ══════════════════════════════════════════════════════════════
# 2. SYNTHESE PAR EQUIPE
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. Synthese par equipe', level=1)

# --- 2.1 PRODUIT ---
doc.add_heading('2.1 Team Produit & Strategie', level=2)

doc.add_heading('Product Manager (LEAD)', level=3)
bullet(doc, 'GO conditionnel — fake CTA d\'abord pour mesurer l\'appetence', 'Avis : ')
bullet(doc, "Proposition de valeur forte : passer du declaratif approximatif au factuel ancre dans les contrats reels")
bullet(doc, "Upload simple + analyse automatique = experience utilisateur fluide")
bullet(doc, "KPI cle : taux de clic sur CTA \"Verifier ma couverture\" > 15%")
bullet(doc, "Risque principal : qualite d'extraction OCR sur documents varies", 'Risque : ')

doc.add_heading('Sales Architect', level=3)
bullet(doc, 'GO — forte valeur commerciale', 'Avis : ')
bullet(doc, "Le gap documente (besoin vs couverture) est un levier de conversion majeur")
bullet(doc, "Formation conseillers necessaire (lecture resultats adequation + posture non-agressive)")
bullet(doc, "Wording neutre et factuel, pas de recommandation produit Baloise directe", 'Condition : ')

doc.add_heading('Decision Scientist', level=3)
bullet(doc, 'GO — moteur d\'adequation techniquement faisable', 'Avis : ')
bullet(doc, "3 couches : extraction automatique (OCR) → matching engine → generation d'actions")
bullet(doc, "Score de confiance : high >= 0.75, medium 0.45-0.74, low < 0.45")
bullet(doc, "Statuts categoriques par garantie : couvert / partiel / non couvert / non evaluable")
bullet(doc, "Eviter la fausse precision — pas de score d'adequation a 73.2%", 'Alerte : ')

# --- 2.2 GOUVERNANCE ---
doc.add_heading('2.2 Team Gouvernance & Conformite', level=2)

doc.add_heading('Compliance Officer', level=3)
bullet(doc, 'GO SOUS CONDITIONS STRICTES', 'Avis : ')
bullet(doc, "Qualification IDD obligatoire : l'outil doit etre qualifie \"information\" (pas \"conseil\")")
bullet(doc, "Disclaimer obligatoire : \"Cette analyse ne constitue pas un conseil en assurance\"")
bullet(doc, "Consentement explicite RGPD avant tout upload de contrats")
bullet(doc, "Si donnees de sante (prevoyance) : consentement art. 9 separe")
bullet(doc, "Sans qualification IDD, le projet ne peut pas demarrer", 'BLOQUANT : ')

doc.add_heading('Legal Counsel', level=3)
bullet(doc, 'GO sous conditions suspensives', 'Avis : ')
bullet(doc, "Responsabilite : si l'outil dit \"vous etes couvert\" alors que non → risque contentieux")
bullet(doc, "Propriete intellectuelle : contrats uploades = propriete du client, pas d'exploitation")
bullet(doc, "Secret des affaires : donnees concurrentielles non exploitables a des fins commerciales")
bullet(doc, "Faire valider le disclaimer par le service juridique Baloise Luxembourg", 'Condition : ')

doc.add_heading('Risk Manager', level=3)
bullet(doc, 'GO conditionnel avec 4 gates de validation', 'Avis : ')
bullet(doc, "Risque reputation : faux positif \"couvert\" → sinistre non couvert → blame Baloise")
bullet(doc, "Risque reglementaire : requalification en conseil par le CAA")
bullet(doc, "Risque donnees : fuite de contrats concurrents → sanctions CNPD")
bullet(doc, "Gates : (G1) DPIA, (G2) IDD, (G3) UAT metier, (G4) revue juridique pre-prod")

doc.add_heading('Internal Audit', level=3)
bullet(doc, 'GO conditionnel — maturite actuelle 1/5 pour le perimetre documents', 'Avis : ')
bullet(doc, "Audit trail complet requis (upload, extraction, suppression)")
bullet(doc, "Mise a jour du plan d'audit annuel pour inclure le nouveau perimetre")

# --- 2.3 TECHNIQUE ---
doc.add_heading('2.3 Team Technique Assurance', level=2)

doc.add_heading('Actuaire Senior', level=3)
bullet(doc, 'GO avec reserves methodologiques', 'Avis : ')
bullet(doc, "Adequation DOIT rester categorique (couvert/non/partiel), PAS numerique")
bullet(doc, "Referentiel 30 garanties = estimation basse, prevoir 40-50 pour le marche LU")
bullet(doc, "Frontaliers (45% de l'emploi) = point nevralgique")
bullet(doc, "Ne pas afficher de pourcentage d'adequation — icones couvert/non couvert", 'Reserve : ')

doc.add_heading('Process Architect', level=3)
bullet(doc, 'GO — workflow en 7 phases', 'Avis : ')
bullet(doc, "Upload → Consentement → Extraction auto → Validation utilisateur → Matching → Rapport → Actions")

doc.add_heading('Insurance Product Manager', level=3)
bullet(doc, "Le mapping garanties/quadrants doit etre valide avec l'equipe technique Baloise")
bullet(doc, "Prevoir un statut \"non evaluable\" pour les garanties hors perimetre Baloise")

# --- 2.4 IT ---
doc.add_heading('2.4 Team IT & Securite', level=2)

doc.add_heading('IT Architect', level=3)
bullet(doc, 'GO — pipeline d\'extraction automatique', 'Avis : ')
bullet(doc, "Pipeline : Upload PDF/photo → OCR (Google Document AI EU) → Extraction structuree → Matching → Adequation")
bullet(doc, "Bucket Supabase Storage prive avec RLS par profile_id")
bullet(doc, "Moteur d'adequation TypeScript (fonction pure, testable)")
bullet(doc, "Ne PAS reproduire le dual scoring TS + PL/pgSQL", 'Alerte : ')

doc.add_heading('Security Architect', level=3)
bullet(doc, 'GO sous 3 bloqueurs absolus', 'Avis : ')
bullet(doc, "DPIA obligatoire avant toute implementation")
bullet(doc, "Consentement art. 9 si donnees de sante")
bullet(doc, "Bucket Supabase prive + chiffrement at-rest + validation magic bytes des fichiers")
bullet(doc, "RLS sur toutes les nouvelles tables + mise a jour delete_my_data/export_my_data")

doc.add_heading('QA Expert', level=3)
bullet(doc, 'GO conditionnel — cahier de test complet prepare', 'Avis : ')
bullet(doc, "70 cas de test en 6 domaines : Upload (19), OCR (14), Matching (12), Scoring (11), Restitution (7), E2E (7)")
bullet(doc, "Architecture CI/CD 3 niveaux : mocks → Tesseract local → API reelle nightly")
bullet(doc, "Corpus de 22 documents de test specifies (PDF natifs, scans, photos, HEIC, corrompus)")
bullet(doc, "Phasage QA : V1 PDF natifs Baloise → V2 multi-assureurs → V3 photos/HEIC")

# --- 2.5 OPS ---
doc.add_heading('2.5 Team Ops', level=2)

doc.add_heading('Underwriting Expert', level=3)
bullet(doc, 'GO conditionnel (confiance 85/100)', 'Avis : ')
bullet(doc, "Structure contrats LU en 5 couches : CP, TRG, CG, CS/Avenants, IPID")
bullet(doc, "48 garanties identifiees : DRIVE (12), HOME (16), TRAVEL (7), B-SAFE (9), FUTUR (4)")

p = doc.add_paragraph()
run = p.add_run('Matrice d\'extractibilite :')
run.bold = True

add_table(doc,
    ['Fiabilite', 'Parametres', 'Confiance'],
    [
        ['Fiable', 'Assureur, produit, dates, prime, niveau couverture global', '> 80%'],
        ['Moderee', 'Plafonds, garanties detaillees, franchises', '55-75%'],
        ['Non fiable', 'Exclusions, mode d\'indemnisation, conditions de garantie', '< 40%'],
    ],
    col_widths=[2.5, 9, 2.5]
)

bullet(doc, "Differences structurelles entre assureurs LU (Baloise, Foyer, La Lux, AXA, Lalux)")
bullet(doc, "Piege #1 : sous-assurance detectable en croisant capital assure vs valeur declaree")
bullet(doc, "Piege #2 : contrats multi-risques a decomposer en sous-garanties")
bullet(doc, "Frontaliers : contrats FR/BE/DE non directement comparables → statut \"non evaluable\"")

# --- 2.6 CREATIVE ---
doc.add_heading('2.6 Team Creative', level=2)

doc.add_heading('Art Director', level=3)
bullet(doc, 'GO — direction UX complete', 'Avis : ')
bullet(doc, "Upload drag & drop + photo mobile (simple, rapide)")
bullet(doc, "Feedback progressif pendant extraction (barre de progression par etape)")
bullet(doc, "Resultats : synthese visuelle avant le detail")
bullet(doc, "Codes couleur : vert (couvert), ambre (partiel), rouge (gap), gris (non evaluable)")
bullet(doc, "Page dediee /results/:id/adequation")

# ══════════════════════════════════════════════════════════════
# 3. CONVERGENCES
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. Convergences majeures (consensus 14/14)', level=1)

add_table(doc,
    ['#', 'Point de convergence', 'Agents concordants'],
    [
        ['1', "Upload automatise — l'utilisateur uploade, le systeme analyse", '14/14'],
        ['2', 'Qualification IDD obligatoire avant dev', 'Compliance, Legal, Risk, Audit, Sales'],
        ['3', 'DPIA obligatoire', 'Compliance, Security, Risk, Legal'],
        ['4', 'Adequation categorique (couvert/non/partiel), pas de %', 'Actuaire, Decision Scientist, PM'],
        ['5', 'Disclaimer non-conseil obligatoire', 'Compliance, Legal, Risk, Sales'],
        ['6', 'Referentiel garanties a valider avec le metier (30 = plancher)', 'Actuaire, Underwriting, QA, PM'],
        ['7', 'Frontaliers = point nevralgique', 'Underwriting, Actuaire, PM'],
        ['8', 'Page dediee (pas surcharger ResultsPage)', 'QA, Art Director, Process Architect'],
        ['9', 'Moteur adequation TypeScript unique', 'IT Architect, QA'],
        ['10', 'Mise a jour delete_my_data / export_my_data', 'Security, QA, Compliance'],
    ],
    col_widths=[1, 9, 5]
)

# ══════════════════════════════════════════════════════════════
# 4. DIVERGENCES
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. Divergences et arbitrages necessaires', level=1)

add_table(doc,
    ['Sujet', 'Position A', 'Position B', 'Arbitrage recommande'],
    [
        ['Score numerique vs categorique',
         'Decision Scientist : score de confiance chiffre (0-1)',
         'Actuaire : categorique uniquement',
         'Categorique pour l\'utilisateur, score interne pour le conseiller'],
        ['CTA factice vs dev direct',
         'Product Manager : fake CTA pour mesurer',
         'Process Architect : dev direct 16 sem.',
         'Fake CTA d\'abord (2j dev, 4 sem. mesure)'],
        ['Referentiel ouvert vs ferme',
         'QA : option "Autre" avec champ libre',
         'Underwriting : referentiel ferme',
         'Ferme + backlog d\'enrichissement'],
        ['Service OCR',
         'IT Architect : Google Document AI EU',
         'QA : Tesseract.js local possible',
         'Google Document AI EU (qualite FR/DE superieure)'],
    ],
    col_widths=[3, 4, 4, 4.5]
)

# ══════════════════════════════════════════════════════════════
# 5. BENCHMARK EXTRACTION AUTOMATIQUE
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('5. Benchmark Extraction Automatique de Contrats', level=1)

doc.add_paragraph(
    "L'approche retenue est l'extraction automatique : l'utilisateur uploade ses contrats "
    "(PDF ou photo), le systeme se charge de tout. Pour selectionner la solution technologique "
    "d'extraction, un benchmark a ete realise par 7 agents specialises (IT Architect, Security Architect, "
    "Compliance Officer, Underwriting Expert, QA Expert, Decision Scientist, Product Manager). "
    "8 solutions ont ete evaluees sur 10+ criteres ponderes."
)

# 5.1 Solutions evaluees
doc.add_heading('5.1 Solutions evaluees', level=2)

add_table(doc,
    ['Solution', 'Type', 'Hebergement EU', 'Cout 500p/mois', 'Integration'],
    [
        ['Mistral AI (Pixtral Large)', 'LLM Vision', 'Natif (Paris)', '15-25 EUR', '4-5 jours'],
        ['Google Document AI', 'OCR structure', 'Region EU dispo', '7-15 EUR', '8-12 jours'],
        ['AWS Textract', 'OCR structure', 'Region EU dispo', '7-10 EUR', '7-10 jours'],
        ['Azure Doc Intelligence', 'OCR structure', 'Region EU dispo', '50-75 EUR', '6-9 jours'],
        ['GPT-4 Vision (OpenAI)', 'LLM Vision', 'PAS EU — ELIMINE', '40-80 EUR', '3-4 jours'],
        ['Claude/Bedrock (Anthropic)', 'LLM Vision', 'EU via Bedrock', '50-100 EUR', '5-7 jours'],
        ['Tesseract.js', 'OCR local', 'Local (zero cloud)', '0 EUR', '3-4 jours'],
        ['Hybride Tesseract + LLM', 'OCR + LLM', 'Mixte', '8-12 EUR', '6-8 jours'],
    ],
    col_widths=[4, 2.5, 3, 2.5, 2.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('GPT-4 Vision est elimine : ')
run.bold = True
run.font.color.rgb = RED
p.add_run(
    "transfert US de donnees de contrats d'assurance incompatible avec les exigences RGPD/CNPD "
    "pour Baloise Luxembourg. Non negociable."
)

# 5.2 Scoring consolide
doc.add_heading('5.2 Scoring consolide', level=2)

p = doc.add_paragraph()
run = p.add_run('Scoring IT Architect (technique, /125) :')
run.bold = True

add_table(doc,
    ['Solution', 'RGPD EU [x3]', 'Qualite photo [x3]', 'Comprehension metier [x3]', 'Cout [x2]', 'Simplicite [x2]', 'SCORE /125', 'Rang'],
    [
        ['Mistral AI', '5', '4', '4', '5', '4', '107', '1er'],
        ['Google DocAI', '4', '5', '2', '5', '2', '90', '2e'],
        ['Azure DocAI', '4', '4', '2', '3', '3', '82', '3e'],
        ['Hybride Tess+LLM', '5', '2', '3', '5', '2', '82', '3e'],
        ['Claude/Bedrock', '4', '3', '5', '2', '3', '77', '5e'],
        ['AWS Textract', '4', '3', '1', '5', '2', '77', '5e'],
        ['Tesseract.js', '5', '1', '1', '5', '3', '72', '7e'],
    ],
    col_widths=[3, 1.8, 2, 2.5, 1.5, 2, 1.5, 1]
)

p = doc.add_paragraph()
run = p.add_run('Scoring PM (strategique, /5 — criteres ponderes C1 a C7) :')
run.bold = True

add_table(doc,
    ['Solution', 'C1 Extraction (25%)', 'C2 RGPD (25%)', 'C3 Integration (15%)', 'C4 Testabilite (12%)', 'Score /5', 'Rang PM'],
    [
        ['Claude/Bedrock', '5', '4', '4', '4', '4.22', '1er'],
        ['Mistral AI', '4', '5', '4', '3', '4.12', '2e'],
        ['Azure DocAI', '4', '4', '3', '4', '3.72', '3e'],
        ['AWS Textract', '3', '4', '3', '4', '3.35', '4e'],
        ['Google DocAI', '3', '4', '3', '4', '3.22', '5e'],
        ['Tesseract.js', '2', '5', '4', '2', '3.22', '6e'],
        ['Hybride', '3', '4', '3', '2', '3.12', '7e'],
    ],
    col_widths=[3, 2.5, 2, 2.5, 2.5, 1.5, 1.5]
)

p = doc.add_paragraph()
run = p.add_run('Observations cles :')
run.bold = True
bullet(doc, "Les OCR purs (Google, AWS, Azure) ne suffisent pas seuls pour l'extraction semantique de contrats d'assurance. Ils necessitent un LLM en aval.")
bullet(doc, "Claude/Bedrock arrive 1er au scoring PM (qualite), Mistral 1er au scoring IT (simplicite/cout).")
bullet(doc, "Tesseract.js seul est insuffisant pour un usage production.")

# 5.3 Divergences inter-agents
doc.add_heading('5.3 Divergences inter-agents sur le choix technologique', level=2)

add_table(doc,
    ['Agent', 'Recommandation', 'Argument principal'],
    [
        ['IT Architect', 'Mistral AI (1er, 107/125)', 'Simplicite, cout, EU natif, pas de surpoids archi, 4-5j integration'],
        ['Security Architect', 'Azure Document AI', 'Pas de prompt injection, deterministe, pas de risque hallucination LLM'],
        ['Compliance Officer', 'OCR structure prefere', 'LLM = charge reglementaire AI Act + IDD plus lourde qu\'un OCR deterministe'],
        ['Underwriting Expert', 'Hybride Azure + LLM', 'Ni l\'OCR ni le LLM seul ne suffit (extractibilite garanties 55-75%)'],
        ['QA Expert', 'Google/Azure (81/100 testabilite)', 'Non-determinisme LLM problematique pour CI/CD, 223 cas de test requis'],
        ['Decision Scientist', 'Architecture agnostique', 'Interface IExtractionService injectable, swap transparent de providers'],
        ['PM (LEAD)', 'Scenario A : Mistral + Claude/Bedrock', 'Comprehension > OCR, bi-provider elimine le vendor lock-in, cout maitrise'],
    ],
    col_widths=[3, 4, 8]
)

p = doc.add_paragraph()
run = p.add_run('Synthese des divergences :')
run.bold = True
doc.add_paragraph(
    "L'IT Architect privilegie la simplicite (Mistral seul). Le Security Architect privilegiait Azure "
    "pour son determinisme. La Compliance prefer un OCR classique (moins de charge reglementaire AI Act). "
    "Le QA Expert veut du deterministe pour les tests CI/CD. Le PM arbitre : le probleme est un probleme "
    "de COMPREHENSION, pas d'OCR. Seul un LLM comprend \"franchise\", \"plafond\", \"garantie implicite\". "
    "L'architecture bi-provider (Mistral principal + Claude/Bedrock fallback) concilie qualite et souverainete."
)

# 5.4 Recommandation PM
doc.add_heading('5.4 Recommandation PM : scenario retenu', level=2)

p = doc.add_paragraph()
run = p.add_run('SCENARIO A — Full LLM bi-provider : Mistral (principal) + Claude/Bedrock (fallback)')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = GREEN

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('5 convictions produit :')
run.bold = True

bullet(doc, "L'extraction de contrats d'assurance est un probleme de COMPREHENSION, pas d'OCR.")
bullet(doc, "La souverainete EU est un critere bloquant, pas un critere de confort (GPT-4V elimine).")
bullet(doc, "La simplicite architecturale est un avantage pour une equipe contrainte (1 dev).")
bullet(doc, "Le cout n'est pas discriminant a 500-5 000 pages/mois (ecart = 15-70 EUR/mois).")
bullet(doc, "L'architecture bi-provider elimine le vendor lock-in (switch = 1 ligne de config).")

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Pourquoi Mistral en principal :')
run.bold = True
bullet(doc, "Societe francaise, serveurs Paris — souverainete EU maximale")
bullet(doc, "Meilleur modele disponible pour le francais")
bullet(doc, "Cout imbattable (15 EUR/mois pour 500 pages)")
bullet(doc, "API compatible format OpenAI, pas de lock-in")
bullet(doc, "Vision native via Pixtral pour scans/photos")

p = doc.add_paragraph()
run = p.add_run('Pourquoi Claude/Bedrock en fallback :')
run.bold = True
bullet(doc, "Meilleure comprehension documentaire du marche (1er au scoring PM)")
bullet(doc, "EU via AWS Bedrock (Frankfurt)")
bullet(doc, "Robustesse en fallback si Mistral indisponible ou insuffisant")

# 5.5 Architecture cible
doc.add_heading('5.5 Architecture cible', level=2)

doc.add_paragraph(
    "Pipeline d'extraction en 9 etapes :"
)

steps = [
    "1. L'utilisateur clique 'Ajouter un contrat' et selectionne un fichier (PDF, JPG, PNG, < 10 MB)",
    "2. Upload vers Supabase Storage (bucket prive 'contracts', RLS par profile_id)",
    "3. INSERT dans table contract_uploads (status: 'pending')",
    "4. Webhook DB declenche une Netlify Background Function (timeout 15 min)",
    "5. Detection du type de document : PDF natif (pdf-parse) vs scan/image (base64)",
    "6. Appel API Mistral : Mistral Large (texte) ou Pixtral Large (image)",
    "7. Validation du JSON retourne (schema Zod, coherence metier)",
    "8. UPDATE contract_uploads (status: 'completed', extracted_data: JSONB)",
    "9. Frontend affiche les resultats, l'utilisateur confirme ou corrige",
]
for step in steps:
    bullet(doc, step)

doc.add_paragraph()

add_table(doc,
    ['Composant technique', 'Detail'],
    [
        ['Frontend', 'React 19 — upload drag & drop + poll/Realtime pour le statut'],
        ['Storage', 'Supabase Storage (bucket prive) — RLS + validation MIME type + 10 MB max'],
        ['Backend', 'Netlify Background Function (Node 18, ESM) — pas de SDK Mistral, simple fetch'],
        ['Extraction texte', 'pdf-parse (PDF natif) — skip si OCR necessaire'],
        ['Extraction image', 'Pixtral Large via API Mistral (vision native)'],
        ['Fallback', 'Claude via AWS Bedrock (Frankfurt) si Mistral indisponible'],
        ['Validation', 'Schema Zod strict — JSON invalide → status \'review\''],
        ['DB', 'Table contract_uploads (JSONB extracted_data, confidence_score, audit trail)'],
        ['Securite', 'Cle API Mistral en env Netlify, CSP inchangee, rate limit 10/jour/user'],
        ['Adequation', 'Moteur TypeScript pur : computeAdequacy(diagnostic, extractedContracts)'],
    ],
    col_widths=[3.5, 11.5]
)

# 5.6 Roadmap extraction
doc.add_heading('5.6 Roadmap extraction', level=2)

add_table(doc,
    ['Phase', 'Periode', 'Scope', 'Go/No-Go'],
    [
        ['Phase 1 : Saisie manuelle', 'T3 2026', 'Wizard React, pas d\'extraction. Construit le corpus de validation (50+ contrats).', 'Qualification IDD + DPIA validee'],
        ['Phase 2 : PDF natifs via Mistral', 'T4 2026', 'Upload PDF → Mistral Large → JSON structure → validation utilisateur', 'Precision >= 70% sur corpus Phase 1'],
        ['Phase 3 : Scans/photos via Vision', 'S1 2027', 'Pixtral Large pour images + Claude Vision en fallback', 'Retour utilisateurs Phase 2 + budget autorise'],
    ],
    col_widths=[3.5, 2, 6.5, 3.5]
)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Budget API mensuel autorise : ')
run.bold = True
p.add_run('50 EUR/mois (alerte a 100 EUR). A 500 pages/mois : ~15 EUR (Mistral), ~50 EUR (Claude/Bedrock fallback).')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('KPIs de succes extraction :')
run.bold = True
bullet(doc, "Precision extraction >= 70% (Phase 2 go/no-go)")
bullet(doc, "Taux de correction utilisateur < 30% des champs")
bullet(doc, "Temps d'extraction < 15 secondes par document")
bullet(doc, "Taux d'echec (status 'failed') < 5%")
bullet(doc, "Cout API < 0.05 EUR/page en regime de croisiere")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. PRE-REQUIS (GATES)
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. Pre-requis avant lancement du developpement (Gates)', level=1)

doc.add_heading('Gate 0 : Validation marche (2 semaines)', level=2)
bullet(doc, 'Deployer CTA factice "Verifier ma couverture" sur ResultsPage')
bullet(doc, 'Mesurer le taux de clic pendant 4 semaines')
bullet(doc, 'Seuil GO : > 15% de taux de clic')

doc.add_heading('Gate 1 : Juridique & Conformite (4 semaines en parallele)', level=2)
bullet(doc, 'Qualification IDD par le service juridique Baloise Luxembourg')
bullet(doc, 'DPIA realisee et validee par le DPO')
bullet(doc, 'Disclaimer valide par Legal')
bullet(doc, 'Analyse des implications art. 9 RGPD (donnees de sante prevoyance)')

doc.add_heading('Gate 2 : Specification technique (2 semaines)', level=2)
bullet(doc, 'Referentiel des garanties normalisees valide par le metier')
bullet(doc, 'Schema DB finalise (migration SQL)')
bullet(doc, 'Pipeline OCR : choix du service, integration, tests de qualite')
bullet(doc, 'Maquettes UX validees (upload simple + resultats)')

doc.add_heading('Gate 3 : Developpement (8 semaines)', level=2)
bullet(doc, 'Pipeline upload : drag & drop + photo mobile → Supabase Storage prive')
bullet(doc, 'Pipeline extraction : OCR → parsing → normalisation → structure')
bullet(doc, 'Matching automatique : garanties extraites → referentiel normalise → quadrants')
bullet(doc, 'Moteur d\'adequation : besoins diagnostiques vs couverture extraite')
bullet(doc, 'Page resultats adequation (synthese visuelle + detail)')
bullet(doc, 'Ecran de validation legere (l\'utilisateur confirme ou corrige les donnees extraites)')
bullet(doc, 'Mise a jour delete_my_data / export_my_data + audit trail')
bullet(doc, 'Vue conseiller des contrats clients')

doc.add_heading('Gate 4 : Recette (2 semaines)', level=2)
bullet(doc, '70 cas de test QA executes (0 bloquant ouvert)')
bullet(doc, 'Suite de non-regression existante 100% verte')
bullet(doc, 'RLS + Storage policies verifiees')
bullet(doc, 'Test de coherence scoring/adequation sur 5 profils types')
bullet(doc, 'Revue securite pre-prod')

# ══════════════════════════════════════════════════════════════
# 7. ESTIMATION
# ══════════════════════════════════════════════════════════════
doc.add_heading('7. Estimation de charge', level=1)

add_table(doc,
    ['Poste', 'Estimation', 'Commentaire'],
    [
        ['Gate 0 (CTA factice)', '2 jours dev + 4 sem. mesure', 'Peut demarrer immediatement'],
        ['Gate 1 (Juridique)', '4 semaines', 'En parallele de Gate 0'],
        ['Gate 2 (Specs)', '2 semaines', 'Apres Gates 0+1 validees'],
        ['Gate 3 (Dev)', '8 semaines', '1 dev (pipeline OCR = +2 sem. vs saisie manuelle)'],
        ['Gate 4 (Recette)', '2 semaines', 'Chevauchement possible avec fin Gate 3'],
        ['TOTAL', '~12 semaines de dev', '+ 4 semaines validation marche en amont'],
    ],
    col_widths=[4, 4, 7]
)

# ══════════════════════════════════════════════════════════════
# 8. RISQUES
# ══════════════════════════════════════════════════════════════
doc.add_heading('8. Risques residuels post-implementation', level=1)

add_table(doc,
    ['Risque', 'Probabilite', 'Impact', 'Mitigation'],
    [
        ['Qualite OCR insuffisante sur photos smartphone', 'Elevee', 'Majeur',
         'Score de confiance + correction manuelle + message "prenez une photo nette"'],
        ['Faux positif "couvert" sans connaitre franchises/exclusions', 'Moyenne', 'Critique',
         'Disclaimer + statut "sous reserve" + formulation prudente'],
        ['Requalification en "conseil" par le CAA', 'Faible', 'Bloquant',
         'Qualification IDD prealable + disclaimer'],
        ['Contrats anciens/scans illisibles', 'Elevee', 'Significatif',
         'Detection qualite + message "document illisible, re-uploadez"'],
        ['Frontaliers mal geres', 'Elevee', 'Significatif',
         'Statut "non evaluable" pour garanties etrangeres'],
        ['Desalignement scoring / adequation', 'Moyenne', 'Majeur',
         'Moteur unique TypeScript + tests de coherence'],
        ['Terminologie inter-assureurs non reconnue', 'Elevee', 'Significatif',
         'Dictionnaire de mapping enrichi iterativement'],
        ['Cout API OCR en production', 'Moyenne', 'Significatif',
         'Budget mensuel + fallback Tesseract local si quota depasse'],
    ],
    col_widths=[4, 2, 2, 7]
)

# ══════════════════════════════════════════════════════════════
# 9. RECOMMANDATION FINALE CONSOLIDEE
# ══════════════════════════════════════════════════════════════
doc.add_heading('9. Recommandation finale consolidee', level=1)

p = doc.add_paragraph()
run = p.add_run('GO CONDITIONNEL — approche en gates')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = GREEN

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Etape immediate (cette semaine) :')
run.bold = True
bullet(doc, 'Deployer le CTA factice sur ResultsPage pour mesurer l\'appetence')
bullet(doc, 'Lancer en parallele la qualification IDD avec le juridique Baloise')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Decision finale apres 4 semaines :')
run.bold = True
bullet(doc, 'Taux de clic CTA > 15% ET qualification IDD favorable → GO developpement')
bullet(doc, 'Taux de clic < 10% → reconsiderer la priorite')
bullet(doc, 'Qualification IDD = "conseil" → revoir le perimetre')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Principes directeurs :')
run.bold = True
bullet(doc, "Simplicite maximale cote utilisateur : upload et c'est tout")
bullet(doc, "Extraction automatique intelligente via LLM (Mistral principal + Claude/Bedrock fallback)")
bullet(doc, "Validation legere : l'utilisateur confirme les donnees extraites en 1 clic")
bullet(doc, "Adequation visuelle : couvert (vert) / partiel (ambre) / gap (rouge)")
bullet(doc, "Disclaimer systematique : analyse indicative, pas un conseil")

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Decision technologique (benchmark) :')
run.bold = True
bullet(doc, "Scenario A retenu — Full LLM bi-provider (Mistral + Claude/Bedrock)")
bullet(doc, "Provider principal : Mistral AI (souverainete EU, cout, qualite FR)")
bullet(doc, "Provider fallback : Claude/Bedrock (qualite extraction, fiabilite)")
bullet(doc, "Phase 2 conditionne a : precision >= 70% sur corpus Phase 1 + DPIA validee")
bullet(doc, "Budget API autorise : 50 EUR/mois (alerte a 100 EUR)")

doc.add_paragraph()
doc.add_paragraph(
    "La fonctionnalite a une forte valeur strategique (differenciation marche, conversion, "
    "donnees concurrentielles) mais les pre-requis reglementaires sont non negociables. "
    "L'approche en gates minimise le risque d'investissement inutile. "
    "La selection de Mistral + Claude/Bedrock concilie souverainete, qualite et simplicite architecturale."
)

# ══════════════════════════════════════════════════════════════
# ANNEXE A
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Annexe A : Standards du marche luxembourgeois', level=1)

doc.add_heading('AUTO (DRIVE)', level=2)
add_table(doc,
    ['Garantie', 'Standard marche', 'Premium', 'Insuffisant'],
    [
        ['RC', 'Illimitee corporel (norme LU)', 'Idem', 'Pas de RC = illegal'],
        ['Omnium', 'Vehicule < 5 ans = norme', 'Valeur a neuf 36 mois + protection bonus', 'Vehicule < 3 ans en RC seule'],
        ['Franchise omnium', '250-500 EUR', '0-150 EUR', '> 1 000 EUR'],
        ['Assistance', 'Incluse en base', '24h, vehicule remplacement equivalent', 'Aucune'],
    ]
)

doc.add_heading('HABITATION (HOME)', level=2)
add_table(doc,
    ['Garantie', 'Standard marche', 'Premium', 'Insuffisant'],
    [
        ['Capital contenu', '20-50k EUR', '> 75k EUR', '< 15k EUR'],
        ['RC VP', '2 500 000 EUR', '5 000 000 EUR+', '< 1 250 000 EUR ou absente'],
        ['Franchise', '250-500 EUR', '150-250 EUR', '> 1 000 EUR'],
        ['Objets de valeur', 'Sous-limite 5-10k EUR', '25k EUR+', 'Aucune'],
    ]
)

doc.add_heading('VOYAGE (TRAVEL)', level=2)
add_table(doc,
    ['Garantie', 'Standard marche', 'Premium', 'Insuffisant'],
    [
        ['Frais medicaux', '150-300k EUR', '> 500k EUR', '< 50k EUR (carte bancaire)'],
        ['Annulation', '2-5k EUR/voyage', '> 10k EUR', '< 1k EUR ou absente'],
        ['Duree sejour', '90 jours', '180 jours', '< 30 jours'],
    ]
)

doc.add_heading('PREVOYANCE (B-SAFE)', level=2)
add_table(doc,
    ['Garantie', 'Standard marche', 'Premium', 'Insuffisant'],
    [
        ['Capital deces', '50-100k EUR', '> 200k EUR (2-3x revenu)', '< 25k EUR'],
        ['Invalidite permanente', '100-200k EUR', '> 300k EUR avec doublement', '< 50k EUR'],
        ['Incapacite', '30-50 EUR/jour', '> 75 EUR/jour, carence 0j', '< 20 EUR/jour ou carence > 30j'],
    ]
)

# ══════════════════════════════════════════════════════════════
# ANNEXE B
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Annexe B : Fichiers d\'implementation cles', level=1)

add_table(doc,
    ['Fichier', 'Action requise'],
    [
        ['src/shared/contracts/guarantees.ts', 'Creer — referentiel normalise NormalizedGuarantee'],
        ['src/shared/contracts/adequacy-engine.ts', 'Creer — fonction pure computeAdequacy()'],
        ['src/shared/contracts/ocr-service.ts', 'Creer — interface IOcrService + implementations'],
        ['src/shared/scoring/engine.ts', 'Pas de modification — adequation = couche supplementaire'],
        ['src/shared/scoring/rules.ts', 'Enrichir — recommandations contextualisees par contrat reel'],
        ['src/lib/api/diagnostics.ts', 'Etendre — CRUD contrats + logAuditEvent nouvelles actions'],
        ['src/pages/client/ResultsPage.tsx', 'CTA "Verifier ma couverture" (~ligne 142)'],
        ['supabase/migrations/019_contracts.sql', 'Creer — tables + RLS + Storage + MAJ delete/export'],
        ['log_audit_event() whitelist', 'Ajouter : contract_uploaded, contract_deleted, adequacy_computed'],
    ],
    col_widths=[6, 9]
)

# ══════════════════════════════════════════════════════════════
# ANNEXE C
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Annexe C : Protocole QA benchmark extraction (synthese)', level=1)

doc.add_paragraph(
    "Le QA Expert a produit un protocole de benchmark complet (223 cas de test) "
    "pour evaluer les solutions d'extraction. Synthese ci-dessous ; "
    "le protocole detaille est dans docs/BENCHMARK-EXTRACTION-PROTOCOLE-QA.md."
)

doc.add_heading('Corpus de test requis (25 documents)', level=2)
add_table(doc,
    ['Categorie', 'Nombre', 'Detail'],
    [
        ['PDF natifs Baloise', '5', 'DRIVE, HOME, TRAVEL, B-SAFE, multi-risques'],
        ['PDF natifs concurrents', '5', 'Foyer, La Luxembourgeoise, AXA LU, Lalux, autre'],
        ['Scans 300 dpi', '5', 'Noir/blanc + couleur + qualite variable'],
        ['Photos smartphone', '5', 'Conditions reelles (eclairage, angle, pli)'],
        ['Documents adversariaux', '5', 'Tronques, multi-pages, bilingues FR/DE, HEIC, corrompus'],
    ],
    col_widths=[4, 2, 9]
)

doc.add_heading('Structure des 223 cas de test', level=2)
add_table(doc,
    ['Domaine', 'Cas de test', 'Exemples'],
    [
        ['Extraction brute (OCR)', '42', 'Precision caracteres, tableaux, multi-colonnes, langues'],
        ['Structuration semantique', '38', 'Identification assureur, garanties, plafonds, franchises'],
        ['Normalisation referentiel', '35', 'Mapping terminologie inter-assureurs, synonymes, variantes'],
        ['Anti-hallucination', '28', 'HALL-T01 a HALL-T06 : donnees inventees, extrapolation abusive'],
        ['Robustesse', '32', 'Documents degrades, formats exotiques, timeout, erreurs API'],
        ['Matching / adequation', '25', 'Coherence extraction → referentiel → scoring adequation'],
        ['Performance / cout', '23', 'Latence, tokens, cout/page, scalabilite'],
    ],
    col_widths=[4, 2, 9]
)

doc.add_heading('Matrice de testabilite par solution', level=2)
add_table(doc,
    ['Solution', 'Determinisme', 'Reproductibilite CI/CD', 'Score testabilite /100'],
    [
        ['Google Document AI', 'Eleve (OCR)', 'Excellente', '81'],
        ['Azure Doc Intelligence', 'Eleve (OCR)', 'Excellente', '81'],
        ['AWS Textract', 'Eleve (OCR)', 'Bonne', '76'],
        ['Mistral AI', 'Faible (LLM)', 'Moyenne (non-deterministe)', '65'],
        ['Claude/Bedrock', 'Faible (LLM)', 'Moyenne', '63'],
        ['Tesseract.js', 'Eleve (local)', 'Excellente', '72'],
    ],
    col_widths=[4, 3, 4, 3.5]
)

p = doc.add_paragraph()
run = p.add_run('Recommandation QA : ')
run.bold = True
p.add_run(
    "Le non-determinisme des LLM est le principal defi qualite. Le protocole prevoit un systeme "
    "de golden tests sur corpus annote + tests de regression nightly + monitoring de la derive "
    "de qualite dans le temps. Les tests anti-hallucination (HALL-T01 a T06) sont critiques."
)

# ══════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    'Document genere a partir de l\'analyse collegiale de 14 agents specialises :\n'
    'Product Manager, IT Architect, Security Architect, Decision Scientist, '
    'Compliance Officer, Legal Counsel, Risk Manager, Internal Audit, '
    'Actuaire Senior, Process Architect, Art Director, Sales Architect, '
    'Underwriting Expert, QA Expert.'
)
run.font.size = Pt(8)
run.font.color.rgb = GREY
run.italic = True

# ── Save ────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), 'SYNTHESE-COLLEGIALE-ADEQUATION-CONTRATS.docx')
doc.save(output_path)
print(f'Word document saved to: {output_path}')
